import openai
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.prompts import PromptTemplate
from langchain.tools.retriever import create_retriever_tool
from langchain.agents import AgentExecutor, create_openai_tools_agent
from langchain_community.tools.tavily_search import TavilySearchResults
from langchain.schema import format_document
from langchain.prompts.prompt import PromptTemplate
import streamlit as st
import tempfile
import base64
from langchain.chains import LLMChain
from pydantic import BaseModel, Field
from langchain_community.document_loaders import UnstructuredWordDocumentLoader
from langchain_community.document_loaders import TextLoader
from langchain.memory import ConversationBufferMemory
import os
from langchain.agents import AgentType, initialize_agent, Tool
from vetting_questions import extracted_questions, modifier_terms, LineList, LineListOutputParser
from langchain.schema import SystemMessage
from docx import Document
from langchain.chains import RetrievalQA
import requests
from langchain.text_splitter import CharacterTextSplitter
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_openai import ChatOpenAI
from langchain.retrievers.multi_query import MultiQueryRetriever
from langchain_openai import OpenAIEmbeddings
from typing import List
from langchain.chains import LLMChain, RetrievalQA, MapReduceDocumentsChain, ReduceDocumentsChain, load_summarize_chain
from langchain.chains.combine_documents.stuff import StuffDocumentsChain
from langchain.text_splitter import TokenTextSplitter
from langchain_community.document_loaders import YoutubeLoader
import pandas as pd
from pandasai import Agent
from pandasai import SmartDataframe
from pandasai.llm import OpenAI


# Streamlit UI setup for multi-page application
st.set_page_config(page_title="RAG Demonstration APP", layout="wide", initial_sidebar_state="expanded", page_icon="logo.png")


# Design move app further up and remove top padding
st.markdown('''<style>.css-1egvi7u {margin-top: -4rem;}</style>''',
    unsafe_allow_html=True)
# Design change hyperlink href link color
st.markdown('''<style>.css-znku1x a {color: #9d03fc;}</style>''',
    unsafe_allow_html=True)  # darkmode
st.markdown('''<style>.css-znku1x a {color: #9d03fc;}</style>''',
    unsafe_allow_html=True)  # lightmode
# Design hide "made with streamlit" footer menu area
hide_streamlit_footer = """<style>#MainMenu {visibility: hidden;}
                        footer {visibility: hidden;}</style>"""
st.markdown(hide_streamlit_footer, unsafe_allow_html=True)



openai.api_key = os.environ["OPENAI_API_KEY"]
# Set OpenAI API key
GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
OPENAI_API_KEY = st.secrets["OPENAI_API_KEY"]
GOOGLE_CSE_ID = st.secrets["GOOGLE_CSE_ID"]
# GOOGLE_API_KEY = os.environ["GOOGLE_API_KEY"]
# GOOGLE_CSE_ID = os.environ["GOOGLE_CSE_ID"]
PANDASAI_API_KEY = st.secrets["PANDASAI_API_KEY"]

# # Prompt for the API key
# api_key = getpass.getpass("Enter your Tavily API Key: ")

# # # Set the environment variable
# os.environ["TAVILY_API_KEY"] = api_key

# # Now you can use the environment variable in your application
# print("API Key set successfully!")

api_key = st.secrets["TAVILY_API_KEY"]



# # Setting the environment variable
# os.environ["TAVILY_API_KEY"] = api_key


# Define functions for document processing
# def load_documents(file_path):
#     loader = PyPDFLoader(file_path)
#     return loader.load()

# def split_documents(docs, chunk_size=1000, chunk_overlap=200, add_start_index=True):
#     text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap, add_start_index=add_start_index)
#     return text_splitter.split_documents(docs)

# def embed_and_index(docs):
#     embeddings = OpenAIEmbeddings()
#     vectorstore = FAISS.from_documents(docs, embeddings)
#     return vectorstore

# def format_docs(docs):
#     return "\n\n".join(doc.page_content for doc in docs)

# def split_documents(docs, chunk_size=1000, chunk_overlap=200, add_start_index=True):
#     text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap, add_start_index=add_start_index)
#     return text_splitter.split_documents(docs)

import tempfile

# Define functions for document processing
def load_documents(file_path):
    if file_path.endswith('.pdf'):
        loader = PyPDFLoader(file_path)
    elif file_path.endswith('.docx'):
        loader = UnstructuredWordDocumentLoader(file_path)
    elif file_path.endswith('.txt'):
        loader = TextLoader(file_path)
    else:
        raise ValueError(f"Unsupported file type for {file_path}")
    return loader.load()

def split_documents(docs, chunk_size=1000, chunk_overlap=200, add_start_index=True):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap, add_start_index=add_start_index)
    return text_splitter.split_documents(docs)

def embed_and_index(docs):
    embeddings = OpenAIEmbeddings()
    vectorstore = FAISS.from_documents(docs, embeddings)
    return vectorstore

def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)


# Function to read file content as string
def get_file_content_as_string(file_path):
    with open(file_path, 'rb') as f:
        binary_file_data = f.read()
    return base64.b64encode(binary_file_data).decode('utf-8')


def create_download_link(file_path, file_name):
    file_content = get_file_content_as_string(file_path)
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{file_content}" download="{file_name}">Download the responses</a>'
    return href

def google_search(query):
    try:
        endpoint = "https://www.googleapis.com/customsearch/v1"
        params = {
            "key": GOOGLE_API_KEY,
            "cx": GOOGLE_CSE_ID,
            "q": query
        }
        response = requests.get(endpoint, params=params)
        results = response.json().get("items", [])
        return [result["link"] for result in results]
    except Exception as e:
        st.error(f"Error during web search: {e}")
        return []



import tempfile

# Function to process the uploaded document
@st.cache_resource
def process_document(file_paths):
    all_docs = []
    for file_path in file_paths:
        if file_path.endswith('.pdf'):
            loader = PyPDFLoader(file_path)
        elif file_path.endswith('.docx'):
            loader = UnstructuredWordDocumentLoader(file_path)
        elif file_path.endswith('.txt'):
            loader = TextLoader(file_path)
        else:
            raise ValueError(f"Unsupported file type for {file_path}")

        try:
            pages = loader.load_and_split()
        except Exception as e:
            st.error(f"Error loading document {file_path}: {e}")
            continue

        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
        docs = splitter.split_documents(pages)
        all_docs.extend(docs)
    
    embeddings = OpenAIEmbeddings(model="text-embedding-3-large")
    vectorstore = FAISS.from_documents(all_docs, embeddings)
    return vectorstore

@st.cache_resource
def process_documents(file_path):
    try:
        loader = PyPDFLoader(file_path)
        pages = loader.load_and_split()
        splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
        docs = splitter.split_documents(pages)
        embeddings = OpenAIEmbeddings(model="text-embedding-3-large")
        retriever = FAISS.from_documents(docs, embeddings).as_retriever()
        return retriever
    except Exception as e:
        st.error(f"Error processing document {file_path}: {e}")
        return None

def handle_uploaded_files(uploaded_file):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(uploaded_file.getvalue())
        return tmp.name

# Function to handle the uploaded files and return paths
def handle_uploaded_file(uploaded_files):
    file_paths = []
    for uploaded_file in uploaded_files:
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp:
            tmp.write(uploaded_file.getvalue())
            file_paths.append(tmp.name)
    return file_paths

 #Summarize functions
def theme_summary(docs):
    llm = ChatOpenAI(temperature=0)

    # Map
    map_template = """The following is a set of documents
    {docs}
    Based on this list of docs, please identify the main themes.
    Helpful Answer:"""
    
    map_prompt = PromptTemplate.from_template(map_template)
    map_chain = LLMChain(llm=llm, prompt=map_prompt)

    # Reduce
    reduce_template = """The following is set of summaries:
    {docs}
    Take these docs and distill it into a final, consolidated summary of the document using the main themes as a guide. For each main idea, include essential supporting details that are necessary to understand the context and significance of the main points.
    Keep your summary objective and unbiased.
    Helpful Answer:"""
    reduce_prompt = PromptTemplate.from_template(reduce_template)

    reduce_chain = LLMChain(llm=llm, prompt=reduce_prompt)

    combine_documents_chain = StuffDocumentsChain(
        llm_chain=reduce_chain, document_variable_name="docs"
    )

    reduce_documents_chain = ReduceDocumentsChain(
        combine_documents_chain=combine_documents_chain,
        collapse_documents_chain=combine_documents_chain,
        token_max=4000,
    )

    map_reduce_chain = MapReduceDocumentsChain(
        llm_chain=map_chain,
        reduce_documents_chain=reduce_documents_chain,
        document_variable_name="docs",
        return_intermediate_steps=False,
    )

    text_splitter = CharacterTextSplitter.from_tiktoken_encoder(
        chunk_size=1000, chunk_overlap=0
    )
    split_docs = text_splitter.split_documents(docs)

    return map_reduce_chain.run(split_docs)

def deep_summary(docs):
    llm = ChatOpenAI(temperature=0)

    splitter = TokenTextSplitter(model_name="gpt-3.5-turbo-16k", chunk_size=10000, chunk_overlap=100)
    chunks = splitter.split_documents(docs)

    summarize_chain = load_summarize_chain(llm=llm, chain_type="refine", verbose=True)
    summary = summarize_chain.run(chunks)

    return summary

def save_summary_to_word(summary, file_name):
    doc = Document()
    doc.add_heading('Summary', 0)
    doc.add_paragraph(summary)
    doc_path = f"{file_name}.docx"
    doc.save(doc_path)
    return doc_path

def download_transcript(transcript, file_name):
    doc_path = f"{file_name}.txt"
    with open(doc_path, 'w') as file:
        file.write(transcript)
    return doc_path

def resume_cover_letter_page():
    prompt_template = """
    
    A your task is to produce {output}. Ensure the {output} is {modifier} and informed by the background documents provided {context}, address the central topic{message}, and integrate any additional details as specified {additional_context}.

    """
    PROMPT = PromptTemplate(
        template=prompt_template, 
        input_variables=["chat_history","context", "message", "additional_context", "output"]
    )

    llm = ChatOpenAI()
    
    output_parser = LineListOutputParser()

    QUERY_PROMPT = PromptTemplate(
    input_variables=["question"],
    template="""You are an AI language model assistant. Your task is to generate five 
    different versions of the given user question to retrieve relevant documents from a vector 
    database. By generating multiple perspectives on the user question, your goal is to help
    the user overcome some of the limitations of the distance-based similarity search. 
    Provide these alternative questions separated by newlines.
    Original question: {question}""")
    llm_chain = LLMChain(llm=llm, prompt=QUERY_PROMPT, output_parser=output_parser)

    question = "What are the approaches to Task Decomposition?"

    # Class for Prompt Input
    class PromptInput(BaseModel):
        context: str
        message: str
        additional_context: str
        chat_history: str

    # memory = ConversationBufferMemory(memory_key="chat_history")

    # Initialize the LLM and LLMChain 
    
    chain = LLMChain(llm=llm, prompt=PROMPT, verbose=True)

    # Streamlit UI setup
    st.subheader("VectorDB Document Generator")

    # Upload PDF and process it
    uploaded_file = st.file_uploader("Upload Documents (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"], accept_multiple_files=True)
    if uploaded_file:
        file_path = handle_uploaded_file(uploaded_file)
        retriever = process_document(file_path)
        st.success("Document processed. Please enter additional context and custom instructions.")

        # Input fields for job description and custom instructions
        with st.expander("Information Input Section", expanded=True):
            message = st.text_area("Message", "Input specific instructions, tasks or directions expected within the generatated output")
            additional_context = st.text_area("Custom instructions", "Enter any specific instructions. If providing ")
            output = st.text_input("Output Type", "Enter expected output e.g. blog post, website content, infographic ")
            modifier = st.selectbox("Choose a predefined query:", modifier_terms)

        
            input_container = st.container()
            with input_container:
                temperature = st.slider("Adjust chatbot specificity:", min_value=0.0, max_value=1.0, value=0.5, step=0.05)
                llm.temperature = temperature
            
            input_container = st.container()
            with input_container:
                chat_model = st.selectbox('What model would you live to choose',('gpt-3.5-turbo-0125','gpt-4o'))
                llm.model_name = chat_model
            
            # col1, col2, col3 = st.columns([5, 5, 5])
            
            # with col1:
            #     num_to_retrieve = st.selectbox('how many documents to retrieve',('5', '10','15','20'))
            
            # with col2:
            #     st.write("\n")  # add spacing

            # with col3:
            #     doc_to_generate = st.selectbox('how many documents to generate',('','3', '4','5','7'))
            
            # input_container = st.container()
            # with input_container:
            #     num_to_retrieve = st.selectbox('how many documents to retrieve',('5', '10','15','20'))

            # input_container = st.container()
            # with input_container:
            #     doc_to_generate = st.selectbox('how many documents to generate',('','3', '4','5','7'))

            if st.button("Generate Document"):
                with st.spinner('Generating your document...'):
                    # retriever_from_llm = MultiQueryRetriever.from_llm(retriever=retriever.as_retriever(), llm=llm)
                    # retriever_from_llm = MultiQueryRetriever(retriever=retriever.as_retriever(), llm_chain=llm_chain, parser_key="lines")
                    # docs = retriever_from_llm.get_relevant_documents(query=message)
                    docs = retriever.similarity_search(query=message, k=4)
                    inputs = [{"context": doc.page_content, "message": message, "additional_context": additional_context, "output": output, "modifier": modifier} for doc in docs]
                    results = chain.apply(inputs)
                    text_results = []
                    for d in results:
                        text = d["text"]
                        formatted_text = f"{text}"
                        if formatted_text not in text_results:
                            text_results.append(formatted_text)

                    st.write("Generated Document", " ".join(text_results))
            else:
                st.write("Please upload a Document to start.")


# Modified document_search_retrieval_page function
def document_search_retrieval_page():
    st.subheader("Document Search and Retrieval")

    uploaded_file = st.file_uploader("Upload Documents (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"], accept_multiple_files=True)
    if uploaded_file:
        file_paths = handle_uploaded_file(uploaded_file)
        vectorstore = process_document(file_paths)
        if not vectorstore:
            return

        retriever = vectorstore.as_retriever(search_type="similarity", search_kwargs={"k": 6})

        pdf_retriever_tool = create_retriever_tool(
            retriever,
            "document_question_and_answer_and_generation",
            "Useful for when you need to answer questions about the uploaded document. Input should be a search query or a given action for information retrieval and generation. Please answer thoroughly and thoughtfully."
        )

        memory = ConversationBufferMemory()

        tavily_search_tool = TavilySearchResults()
        llm = ChatOpenAI()
        tools = [pdf_retriever_tool, tavily_search_tool]
        agent_prompt = ChatPromptTemplate.from_messages(
            [
                ("system", "You are an intelligent Assistant that offers clear and concise answers to queries. You need to be very thorough. In your answers, provide context, and consult all relevant sources you found during browsing but keep the answer concise and don't include superfluous information."),
                MessagesPlaceholder("chat_history", optional=True),
                ("human", "{input}"),
                MessagesPlaceholder("agent_scratchpad"),
            ]
        )
        agent = create_openai_tools_agent(llm, tools, agent_prompt)
        agent_executor = AgentExecutor(agent=agent, tools=tools, memory=memory)

        input_container = st.container()
        with input_container:
            temperature = st.slider("Adjust chatbot specificity:", min_value=0.0, max_value=1.0, value=0.5, step=0.05)
            llm.temperature = temperature

        input_container = st.container()
        with input_container:
            chat_model = st.selectbox('What model would you like to choose', ('gpt-3.5-turbo-0125', 'gpt-4'))
            llm.model_name = chat_model

        chat_container = st.container()
        with chat_container:
            if "messages" not in st.session_state:
                st.session_state.messages = []
            for message in st.session_state.messages:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])

            user_input = st.text_input("Ask a question about the uploaded document:")

            if st.button('Query Document') and user_input:
                with st.spinner('Processing your question...'):
                    try:
                        response = agent_executor.invoke({"input": user_input})
                        st.session_state.messages.append({"role": "assistant", "content": response["output"]})
                        with st.chat_message("assistant"):
                            st.markdown(response["output"])
                    except Exception as e:
                        st.error(f"An error occurred: {e}")

            if st.button('Search Web') and user_input:
                with st.spinner('Searching the web...'):
                    try:
                        tavily_search_result = agent_executor.invoke({
                            "input": user_input
                        })
                        st.session_state.messages.append({"role": "assistant", "content": tavily_search_result["output"]})
                        with st.chat_message("assistant"):
                            st.markdown(tavily_search_result["output"])
                    except Exception as e:
                        st.error(f"An error occurred during web search: {e}")

            # Add the new buttons for summarization
            if st.button('Theme Summary'):
                with st.spinner('Generating theme summary...'):
                    try:
                        docs = []
                        for file_path in file_paths:
                            docs.extend(load_documents(file_path))
                        summary = theme_summary(docs)
                        doc_path = save_summary_to_word(summary, "theme_summary")
                        st.markdown(create_download_link(doc_path, "theme_summary.docx"), unsafe_allow_html=True)
                    except Exception as e:
                        st.error(f"An error occurred: {e}")

            if st.button('Deep Summary'):
                with st.spinner('Generating deep summary...'):
                    try:
                        docs = []
                        for file_path in file_paths:
                            docs.extend(load_documents(file_path))
                        summary = deep_summary(docs)
                        doc_path = save_summary_to_word(summary, "deep_summary")
                        st.markdown(create_download_link(doc_path, "deep_summary.docx"), unsafe_allow_html=True)
                    except Exception as e:
                        st.error(f"An error occurred: {e}")

    else:
        st.write("Please upload a document to start.")

    st.subheader("YouTube Video Transcription and Summary")

    youtube_link = st.text_input("Enter YouTube Video Link:")

    if st.button('Download Transcript'):
        with st.spinner('Downloading transcript...'):
            try:
                loader = YoutubeLoader.from_youtube_url(youtube_link, language=["en", "en-US"])
                transcript = loader.load()
                
                # Join the transcript text
                all_text = "\n".join([chunk.page_content for chunk in transcript])
                doc_path = download_transcript(all_text, "youtube_transcript")
                
                st.markdown(create_download_link(doc_path, "youtube_transcript.txt"), unsafe_allow_html=True)
            except Exception as e:
                st.error(f"An error occurred: {e}")

    if st.button('Summarize Video Content'):
        with st.spinner('Summarizing video content...'):
            try:
                loader = YoutubeLoader.from_youtube_url(youtube_link, language=["en", "en-US"])
                transcript = loader.load()

                splitter = TokenTextSplitter(model_name="gpt-3.5-turbo-16k", chunk_size=10000, chunk_overlap=100)
                chunks = splitter.split_documents(transcript)
                
                summary = deep_summary(chunks)
                doc_path = save_summary_to_word(summary, "video_summary")
                
                st.markdown(create_download_link(doc_path, "video_summary.docx"), unsafe_allow_html=True)
            except Exception as e:
                st.error(f"An error occurred: {e}")
    
        st.subheader("Ask Questions About Your CSV/Excel Data")

    csv_file = st.file_uploader("Upload CSV/Excel File", type=["csv", "xlsx"])
    
    if csv_file:
        df = None
        try:
            if csv_file.name.endswith('.csv'):
                df = pd.read_csv(csv_file)
            elif csv_file.name.endswith('.xlsx'):
                df = pd.read_excel(csv_file)
        except Exception as e:
            st.error(f"Error loading file: {e}")

        if df is not None:
            st.write("Data Preview:", df.head())

            # Integrate OpenAI LLM with PandasAI
            PANDASAI_API_KEY = st.secrets["PANDASAI_API_KEY"]
            llm = OpenAI(api_token=os.getenv("OPENAI_API_KEY"))  # Uses the API key from environment variable
            pandas_ai = SmartDataframe(df, config={"llm": llm})

            query = st.text_input("Ask a question about the data:")
            if st.button('Query Data') and query:
                with st.spinner('Processing your query...'):
                    try:
                        response = pandas_ai.chat(query)
                        st.write(f"Answer: {response}")
                    except Exception as e:
                        st.error(f"An error occurred: {e}")

    else:
        st.write("Please upload a CSV or Excel file to start.")



def vetting_assistant_page():
    st.title("Vetting Assistant Chatbot")

    if "uploaded_pdf_path" not in st.session_state or "retriever" not in st.session_state:
        uploaded_file = st.file_uploader("Upload a PDF containing the terms of service", type=["pdf"])
        # uploaded_file = st.file_uploader("Upload Documents (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"], accept_multiple_files=True)

        if uploaded_file:
            file_path = handle_uploaded_files(uploaded_file)
            st.session_state.uploaded_pdf_path = file_path
            st.session_state.retriever = process_documents(st.session_state.uploaded_pdf_path)
    else:
        st.write("Using previously uploaded PDF. If you want to use a different PDF, please refresh the page.")

    app_name = st.text_input("Enter the name of the app:")

    if "retriever" in st.session_state:
        llm = ChatOpenAI()
        tools = [
            Tool(
                name="vetting_tool",
                description="Tool for retrieving infomration related to security and privacy. Input should be a search query or a given action for information retrieval and generation.",
                func=RetrievalQA.from_llm(llm=llm, retriever=st.session_state.retriever, return_source_documents=True)
            )
        ]
        # tool = create_retriever_tool(
        #     st.session_state.retriever,
        #     "search_terms_service",
        #     "Searches and returns an application's privacy and data policies and terms of use.",
        # )
        # tools = [tool]
        agent_kwargs = {
            "system_message": SystemMessage(content="You are an intelligent Vetting Assistant, "
                                                    "expertly designed to analyze and extract key "
                                                    "information from terms of service and privacy policy documents. "
                                                    "Your goal is to assist users in understanding "
                                                    "complex legal documents and provide clear, "
                                                    "concise answers to their queries. Use only the retrieved data from the pdf when responding to questions."
                                                    "Include citations including section and page number with your response. ")
        }
        agent = initialize_agent(agent=AgentType.OPENAI_FUNCTIONS, tools=tools, llm=llm, agent_kwargs=agent_kwargs,
                                 verbose=True)
        # agent = create_conversational_retrieval_agent(llm, tools)

        input_container = st.container()
        with input_container:
            temperature = st.slider("Adjust chatbot specificity:", min_value=0.0, max_value=1.0, value=0.5, step=0.05)
            llm.temperature = temperature
        
        input_container = st.container()
        with input_container:
            chat_model = st.selectbox('What model would you like to choose',('gpt-3.5-turbo-0125','gpt-4o'))
            llm.model_name = chat_model

        st.write("Ask any question related to the vetting process:")
        query_option = st.selectbox("Choose a predefined query:", extracted_questions)
        user_input = st.text_input("Your Question:")

        if st.button('Start Vetting') and user_input:
            with st.spinner('Processing your question...'):
                try:
                    response = agent.run(user_input)
                    # response = agent({"input": "user_input"})
                    st.write(f"Answer: {response}")
                except Exception as e:
                    st.error(f"An error occurred: {e}")

        st.write("Note: The chatbot retrieves answers from the uploaded document.")

        if 'running_queries' not in st.session_state:
            st.session_state.running_queries = False

        placeholder_message = f"{app_name} is being vetted for compliance and its policies provided in context. Does {app_name} meet this criteria?"
        all_queries = [f"{question} {placeholder_message}" for question in extracted_questions]

        if st.button('Run All Queries'):
            with st.spinner('Processing all queries...'):
                st.session_state.running_queries = True
                doc = Document()
                doc.add_heading('Vetting Assistant Responses', 0)

                for question in all_queries:
                    if not st.session_state.running_queries:
                        break
                    try:
                        response = agent.run(question)
                        doc.add_heading('Q:', level=1)
                        doc.add_paragraph(question)
                        doc.add_heading('A:', level=1)
                        doc.add_paragraph(response)
                    except Exception as e:
                        doc.add_paragraph(f"Error for question '{question}': {e}")

                doc_path = "vetting_responses.docx"
                doc.save(doc_path)
                st.markdown(create_download_link(doc_path, "vetting_responses.docx"), unsafe_allow_html=True)

        if st.button('Stop Queries'):
            st.session_state.running_queries = False

        if st.button('Search Web'):
            with st.spinner('Searching the web...'):
                links = google_search(user_input)
                st.write("Top search results:")
                for link in links:
                    st.write(link)

# Streamlit UI setup for multi-page application
st.image('img/image.png')
st.title(":blue[Document Processing and Retrieval Application]")
st.markdown('Generate engaging documents and chat over your files based on your direct quries - powered by Artificial Intelligence (OpenAI GPT-4 and GPT-3.5) Implemented by PCM')
        # '[stefanrmmr](https://www.linkedin.com/in/stefanrmmr/) - '
        # 'view project source code on '
        # '[GitHub](https://github.com/stefanrmmr/gpt3_email_generator)'
st.write('\n')  # add spacing
page = st.sidebar.selectbox("Choose a Tool:", ["VectorDB Document Generator", "Document Search and Retrieval", "Vetting Assistant"])

# Load respective pages
if page == "VectorDB Document Generator":
    resume_cover_letter_page()
elif page == "Document Search and Retrieval":
    document_search_retrieval_page()
elif page == "Vetting Assistant":
    vetting_assistant_page()






