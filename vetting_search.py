import streamlit as st
from langchain.agents import AgentType, initialize_agent, Tool
from langchain.document_loaders import PyPDFLoader
from langchain.chat_models import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.vectorstores.faiss import FAISS
import openai
import os
import requests
from vetting_questions import extracted_questions
import uuid
from docx import Document
import base64
from langchain.document_loaders.url import UnstructuredURLLoader


# Set OpenAI API key
openai.api_key = os.environ["OPENAI_API_KEY"]
# GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
# OPENAI_API_KEY = st.secrets["OPENAI_API_KEY"]
# GOOGLE_CSE_ID = st.secrets["GOOGLE_CSE_ID"]
GOOGLE_API_KEY = os.environ["GOOGLE_API_KEY"]
GOOGLE_CSE_ID = os.environ["GOOGLE_CSE_ID"]


def get_file_content_as_string(file_path):
    with open(file_path, 'rb') as f:
        binary_file_data = f.read()
    return base64.b64encode(binary_file_data).decode('utf-8')


def create_download_link(file_path, file_name):
    file_content = get_file_content_as_string(file_path)
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{file_content}" download="{file_name}">Download the responses</a>'
    return href


@st.cache_data
def process_document(file_path):
    loader = PyPDFLoader(file_path)
    pages = loader.load_and_split()
    splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
    docs = splitter.split_documents(pages)
    embeddings = OpenAIEmbeddings()
    retriever = FAISS.from_documents(docs, embeddings).as_retriever()
    return retriever


@st.cache_data
def process_url_content(url):
    try:
        loader = UnstructuredURLLoader(urls=[url], mode="single")
        pages = loader.load_and_split()
        if not pages:
            raise ValueError("No content was extracted from the provided URL.")

        splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
        docs = splitter.split_documents(pages)
        embeddings = OpenAIEmbeddings()
        retriever = FAISS.from_documents(docs, embeddings).as_retriever()
        return retriever
    except Exception as e:
        st.error(f"An error occurred while processing the URL content: {e}")
        return None


def google_search(query):
    try:
        endpoint = "https://www.googleapis.com/customsearch/v1"
        params = {
            "key": GOOGLE_API_KEY,
            "cx": GOOGLE_CSE_ID,
            "q": query
        }
        response = requests.get(endpoint, params=params)
        results = response.json().get("items", [])
        return [result["link"] for result in results]
    except Exception as e:
        st.error(f"Error during web search: {e}")
        return []


def handle_uploaded_file(uploaded_file):
    unique_filename = f"uploaded_terms_{uuid.uuid4()}.pdf"
    file_path = os.path.join("uploaded_documents", unique_filename)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getvalue())
    return file_path


def vetting_assistant_page():
    st.title("Vetting Assistant Chatbot")

    # File uploader for PDFs
    uploaded_file = st.file_uploader("Upload a PDF containing the terms of service", type=["pdf"])

    # Input field for the app name
    app_name = st.text_input("Enter the name of the app:")

    # Input field for URL
    url_input = st.text_input("Enter a URL to extract content:")

    retriever = None

    if uploaded_file:
        # Generate a unique filename using uuid
        unique_filename = f"uploaded_terms_{uuid.uuid4()}.pdf"
        file_path = os.path.join("uploaded_documents", unique_filename)

        # Ensure the directory exists
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        with open(file_path, "wb") as f:
            f.write(uploaded_file.getvalue())

        retriever = process_document(file_path)

    elif url_input:
        retriever = process_url_content(url_input)

    if retriever:
        llm = ChatOpenAI(temperature=0.5, model="gpt-3.5-turbo-16k")
        tools = [
            Tool(
                name="vetting_tool",
                description="Tool for vetting based on document content",
                func=RetrievalQA.from_chain_type(llm=llm, retriever=retriever)
            )
        ]

        agent = initialize_agent(agent=AgentType.OPENAI_FUNCTIONS, tools=tools, llm=llm, verbose=True)

        st.write("Ask any question related to the vetting process:")

        # Dropdown for predefined questions
        query_option = st.selectbox("Choose a predefined query:", extracted_questions)

        # Text input for custom question or modification
        user_input = st.text_input("Your Question:", value=query_option)

        if st.button('Start Vetting') and user_input:
            with st.spinner('Processing your question...'):
                try:
                    response = agent.run(user_input)
                    st.write(f"Answer: {response}")
                except Exception as e:
                    st.error(f"An error occurred: {e}")

        st.write("Note: The chatbot retrieves answers from the uploaded document or provided URL.")

        # Initialize session state variable for running queries
        if 'running_queries' not in st.session_state:
            st.session_state.running_queries = False

        # Placeholder message to be appended to each query
        placeholder_message = f"{app_name} is being vetted for compliance and its policies provided in context. Does {app_name} meet this criteria?"

        # Button to run all queries
        if st.button('Run All Queries'):
            with st.spinner('Processing all queries...'):
                st.session_state.running_queries = True
                doc = Document()
                doc.add_heading('Vetting Assistant Responses', 0)

                # Append the user's custom query and the placeholder message to the list of predefined queries
                all_queries = [f"{question} {placeholder_message}" for question in extracted_questions]

                for question in all_queries:
                    # Check if we should stop running queries
                    if not st.session_state.running_queries:
                        break

                    try:
                        response = agent.run(question)
                        doc.add_heading('Q:', level=1)
                        doc.add_paragraph(question)
                        doc.add_heading('A:', level=1)
                        doc.add_paragraph(response)
                    except Exception as e:
                        doc.add_paragraph(f"Error for question '{question}': {e}")

                # Save the document
                doc_path = "vetting_responses.docx"
                doc.save(doc_path)

                # Provide a download link
                st.markdown(create_download_link(doc_path, "vetting_responses.docx"), unsafe_allow_html=True)

        # Button to stop all queries
        if st.button('Stop Queries'):
            st.session_state.running_queries = False

        if st.button('Search Web'):
            with st.spinner('Searching the web...'):
                links = google_search(user_input)
                st.write("Top search results:")
                for link in links:
                    st.write(link)


def pdf_chatbot_page():
    st.title("PDF-based Chatbot")

    if "uploaded_pdf_path" not in st.session_state:
        uploaded_file = st.file_uploader("Upload a PDF", type=["pdf"])

        if uploaded_file:
            file_path = handle_uploaded_file(uploaded_file)
            st.session_state.uploaded_pdf_path = file_path
            st.session_state.retriever = process_document(st.session_state.uploaded_pdf_path)
    else:
        st.write("Using previously uploaded PDF. If you want to use a different PDF, please refresh the page.")

    if "retriever" in st.session_state:
        llm = ChatOpenAI(temperature=0.5, model="gpt-3.5-turbo-16k")
        tools = [
            Tool(
                name="pdf_tool",
                description="Tool for querying based on document content",
                func=RetrievalQA.from_chain_type(llm=llm, retriever=st.session_state.retriever)
            )
        ]
        agent = initialize_agent(agent=AgentType.OPENAI_FUNCTIONS, tools=tools, llm=llm, verbose=True)

        instructions_container = st.container()
        with instructions_container:
            st.header("Instructions")
            st.write("""
            - This chatbot provides answers based on the content of the uploaded PDF.
            - Type in your question in the chat input below.
            - Adjust the slider to control the specificity of the chatbot's responses.
            """)

        input_container = st.container()
        with input_container:
            temperature = st.slider("Adjust chatbot specificity:", min_value=0.0, max_value=1.0, value=0.5, step=0.05)
            llm.temperature = temperature

        chat_container = st.container()
        with chat_container:
            if "messages" not in st.session_state:
                st.session_state.messages = []
            for message in st.session_state.messages:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])

            user_input = st.text_input("Ask a question about the uploaded PDF:")

            if st.button('Query PDF') and user_input:
                with st.spinner('Processing your question...'):
                    try:
                        response = agent.run(user_input)
                        st.session_state.messages.append({"role": "assistant", "content": response})
                        with st.chat_message("assistant"):
                            st.markdown(response)
                    except Exception as e:
                        st.error(f"An error occurred: {e}")

            if st.button('Search Web'):
                with st.spinner('Searching the web...'):
                    links = google_search(user_input)
                    st.write("Top search results:")
                    for link in links:
                        st.write(link)


# Streamlit UI Configuration
st.set_page_config(page_title="Vetting Assistant Chatbot", layout="wide", initial_sidebar_state="expanded")
page = st.sidebar.selectbox("Choose a Tool:", ["Vetting Assistant", "PDF Chatbot"])

if page == "Vetting Assistant":
    vetting_assistant_page()
elif page == "PDF Chatbot":
    pdf_chatbot_page()
